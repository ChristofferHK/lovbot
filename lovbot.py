from __future__ import print_function


import docx
import pdfkit
import nltk
import sys
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from tika import parser

# creds til matheusportela for simpletable
# https://github.com/matheusportela/simpletable
import simpletable

# import tensorflow as tf


# nltk.download('punkt')
# nltk.download('wordnet')

# Next step:
# || la brukeren søke på mer enn bare nøkkelord, korte ned svar

# || lemmatisere input og ordene i dokumentet, slik at man kan få flere treff
# på ord som har ulike endelser i bruk

# || Hente raw fra andre kilder, html og ikke bare .txt | html, docx, pdf

# || sortere resultatene etter størst value
# med andre ord, det mest relevante resultatet først

# || lage en frontend

# Setter tabellen vi får i html til en pdf
def html2pdf(filename):
    config_path = 'C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe'
    pdfkit.configuration(wkhtmltopdf=config_path)
    pdfkit.from_url('file:///C:/Users/chris/Documents/ordml/static/resources{}.html'.format(filename),
                    '{}.pdf'.format(filename))


# henter alle paragrafene i word-filen og setter det til en helhetlig tekst
def getDocText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


# Sjekker filtypen og oppretter råteksten basert på den
def generateRaw(filename):
    if (filename.endswith(".txt")):
        # Henter .txt
        with open(filename, 'r', errors='ignore') as fin:
            raw = fin.read().lower()
    elif (filename.endswith(".docx")):
        # Henter .docx
        raw = getDocText(filename)
    elif (filename.endswith(".pdf")):
        # Henter .pdf
        raw = parser.from_file(filename)
        raw = raw['content']
    elif (filename.endswith(".html")):
        print()
    else:
        print("Ukjent filtype.")
    return raw


stemmer = nltk.LancasterStemmer()
lemmer = nltk.stem.WordNetLemmatizer()

raw = generateRaw(sys.argv[1])
filename_base = sys.argv[1][0:-4]

sents = nltk.sent_tokenize(raw)
words = nltk.word_tokenize(raw)


def response(input_key):
    # Term Frequency - Inverse Document Frequency (TF-IDF): TF * IDF
    # Term Frequency - antall ganger ordet dukker opp / antall ord
    # Inverse Document Frequency - log(N/n)
    # N: antall dokumenter, n: antall dokumenter ordet har dukket opp i

    # Cosine_Similarity - hvor like dokumenter er, uavhengig av størrelse
    # Brukes blant annet til ansiktsgjenkjenning og intensjonsklassifisering

    sents.append(input_key)
    tfidf = TfidfVectorizer().fit_transform(sents)
    vals = cosine_similarity(tfidf[-1], tfidf)
    matches = []

    for val in vals:
        c = 0
        for v in val:
            if v.any() != 0:
                matches.append(c)
            c = c + 1
    return matches


def user_input():
    flag = True
    print("BERTIL: Mitt navn er Bertil, og jeg vil svare på alle dine spørsmål relatert til personopplysningsloven.\n\nVil du avslutte så skriv stopp eller farvel.")

    while flag:
        print("BERTIL: Hva leter du etter?")
        user_response = input("> ")
        user_response = user_response.lower()

        if (user_response == "stopp" or user_response == "farvel"):
            print("BERTIL: Farvel og lykke til!")
            flag = False
        else:
            results = response(user_response)
            if (len(results) == 2):
                print("BERTIL: Jeg fant ett resultat om {}\n".format(
                    user_response))
            elif (len(results) == 1):
                print("BERTIL: Jeg fant ingen resultater om {}".format(
                    user_response))
            else:
                print("BERTIL: Jeg fant {} resultater om {}\n".format(
                    len(results) - 1, user_response))

            if (len(results) > 1):
                print("BERTIL: Vil du ha dette på tabellformat? Ja/Nei")
                yesno = input()
                if (yesno == "ja"):
                    table_format(user_response, results, filename_base)
                    print("BERTIL: Vil du ha tabellen i en pdf? Ja/Nei")
                    yesno = input()
                    if (yesno == "ja"):
                        html2pdf(filename_base)
                else:
                    for counter, r in enumerate(results):
                        if (len(results) <= 2):
                            print("BERTIL: Viser resultatet\n{}".format(
                                sents[results[counter - 1]]))
                        elif (counter != len(results) - 1):
                            print("BERTIL: Viser nummer {} av {}. Trykk Enter for å vise neste\n{}".format(
                                counter + 1, len(results) - 1, sents[results[counter]]))
                            request = input()
                            if (request == "stopp"):
                                print("BERTIL: Vil du avslutte helt? Ja/Nei")
                                yesno = input()
                                if (yesno == "ja"):
                                    flag = False
                                    break
                                else:
                                    break

            sents.remove(user_response)


def table_format(word, matches, filename):
    # setter opp resultatlista
    result_list = []
    for match in matches:
        result_list.append(sents[match])
    result_list.remove(result_list[-1])

    data = [str(r) for r in result_list]

    # setter opp css
    if __name__ == "__main__":
        css = """
        table.mytable {
            font-family: times;
            font-size:14px;
            color:#000000;
            border-width: 1px;
            border-color: #eeeeee;
            border-collapse: collapse;
            background-color: #ffffff;
            width=100%;
            max-width:550px;
            table-layout:fixed;
        }
        table.mytable th {
            border-width: 1px;
            padding: 8px;
            border-style: solid;
            border-color: #eeeeee;
            background-color: #e6eed6;
            color:#000000;
        }
        table.mytable td {
            border-width: 1px;
            padding: 8px;
            border-style: solid;
            border-color: #eeeeee;
        }
        #code {
            display:inline;
            font-family: courier;
            color: #3d9400;
        }
        #string {
            display:inline;
            font-weight: bold;
        }
        """

    # konverterer lista med resultater til passende html format
    table = simpletable.SimpleTable([data],
                                    header_row=[word],
                                    css_class='mytable')

    # setter opp html-siden
    page = simpletable.HTMLPage()
    page.add_table(table)
    page.css = css
    page.save("{}.html".format(filename))


user_input()
